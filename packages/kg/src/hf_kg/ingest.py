"""Универсальные читатели файлов для ingestion (ТЗ P9: PDF, DOCX, TXT, MD,
CSV, XLSX, JSON). Возвращают простые структуры; превращение их в Fact'ы и
чанки — забота доменного адаптера или ядра.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

import docx
import openpyxl


def read_docx_paragraphs(path: str | Path) -> list[str]:
    d = docx.Document(str(path))
    out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            seen = set()
            for c in row.cells:
                t = c.text.strip()
                if t and t not in seen:
                    out.append(t)
                    seen.add(t)
    return out


def read_text_paragraphs(path: str | Path) -> list[str]:
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    return [p.strip() for p in text.split("\n\n") if p.strip()]


def read_csv_rows(path: str | Path) -> list[dict]:
    """Строки CSV как словари {заголовок: значение}; разделитель определяется."""
    raw = Path(path).read_text(encoding="utf-8-sig", errors="replace")
    delimiter = ";" if raw.count(";") > raw.count(",") else ","
    reader = csv.DictReader(raw.splitlines(), delimiter=delimiter)
    return [dict(r) for r in reader if any((v or "").strip() for v in r.values())]


def read_xlsx_rows(path: str | Path, sheet: str | None = None) -> list[dict]:
    """Строки первого (или указанного) листа как словари по заголовку."""
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.worksheets[0]
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []
    header = [str(h).strip() if h is not None else f"col{i}"
              for i, h in enumerate(rows[0])]
    out = []
    for row in rows[1:]:
        if all(v is None for v in row):
            continue
        out.append({header[i]: row[i] for i in range(min(len(header), len(row)))})
    return out


def read_json_data(path: str | Path):
    return json.loads(Path(path).read_text(encoding="utf-8"))


def read_pdf_text(path: str | Path) -> list[str]:
    """Постраничный текст PDF (best effort; без OCR)."""
    try:
        from pypdf import PdfReader
    except ImportError:  # pypdf опционален: без него PDF даёт пустой список
        return []
    try:
        reader = PdfReader(str(path))
        return [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception:
        return []
