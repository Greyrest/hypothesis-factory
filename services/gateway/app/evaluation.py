"""Leave-one-out оценка: покрытие тем эталонных мозговых штурмов экспертов."""
from __future__ import annotations

import re
from pathlib import Path

import docx

# темы для сравнения с эталоном (грубое покрытие по ключевым словам)
EVAL_TOPICS = {
    "мельниц|футеровк|шаров|измельчен": "измельчение",
    "гидроциклон|классификат|насадок|классификац|грохо": "классификация",
    "флотац|контактн|чан|плотност|фронт": "флотация",
    "реагент|finfix": "реагенты",
    "дробилк|гранулометр|зазор": "дробление",
    "магнитн|сепарац": "магнитная сепарация",
    "хвост.*возврат|возврат.*хвост": "возврат хвостов",
    "автоматизац|автоматическ|контрол": "автоматизация/контроль",
}


def _read_docx_paragraphs(path: Path) -> list[str]:
    d = docx.Document(str(path))
    out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            seen = set()
            for c in row.cells:
                t = c.text.strip()
                if t and t not in seen:
                    out.append(t)
                    seen.add(t)
    return out


def _topics(texts: list[str]) -> set[str]:
    out = set()
    joined = [t.lower() for t in texts]
    for pattern, topic in EVAL_TOPICS.items():
        if any(re.search(pattern, t) for t in joined):
            out.add(topic)
    return out


def evaluate(results: dict[str, dict], data_dir: Path) -> dict:
    """Покрытие тем эталонного мозгового штурма сгенерированными гипотезами."""
    report = {}
    for hyp_file in sorted(data_dir.glob("Пример */Гипотезы*.docx")):
        label = None
        # сопоставляем docx с фабрикой по xlsx в той же папке
        xlsx = next(hyp_file.parent.glob("Хвосты*.xlsx"), None)
        if xlsx:
            label = xlsx.stem.replace("Хвосты", "").replace("_2", "").strip()
        if label not in results:
            continue
        ref_lines = [l for l in _read_docx_paragraphs(hyp_file)
                     if re.match(r"^\d+\.", l)]
        ref_topics = _topics(ref_lines)
        gen_topics = _topics([h["title"] + " " + h["hypothesis"]
                              for h in results[label]["hypotheses"]])
        covered = ref_topics & gen_topics
        report[label] = {
            "expert_topics": sorted(ref_topics),
            "generated_topics": sorted(gen_topics),
            "covered": sorted(covered),
            "coverage_pct": round(100 * len(covered) / len(ref_topics), 0)
            if ref_topics else None,
        }
    return report
