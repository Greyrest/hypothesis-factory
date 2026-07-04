"""Экспорт результатов на лету: JSON (интеграция), CSV (задачи),
Markdown и DOCX (отчёты). Генерируется по запросу GET /export —
на диск ничего не пишется.
"""
from __future__ import annotations

import csv
import io
import json


def _components(result: dict) -> list[dict]:
    """Компоненты результата; для старых снапшотов — из ключей summary."""
    comps = result.get("components")
    if comps:
        return comps
    ids = sorted(result.get("summary", {}).get("losses_t", {}))
    return [{"id": i, "label": i} for i in ids]


def to_json(result: dict) -> str:
    return json.dumps(result, ensure_ascii=False, indent=2)


def to_csv(result: dict) -> str:
    comps = _components(result)
    buf = io.StringIO()
    w = csv.writer(buf, delimiter=";")
    comp_cols = [col for c in comps
                 for col in (f"addressable_{c['id']}_t",
                             f"kpi_delta_{c['id']}_t_min",
                             f"kpi_delta_{c['id']}_t_max")]
    w.writerow(["rank", "priority", "title", "category", "hypothesis",
                *comp_cols,
                "feasibility_1_5", "novelty_1_5", "risk_1_5",
                "status", "streams", "sources"])
    for h in result["hypotheses"]:
        eff = h["expected_effect"]
        comp_vals = []
        for c in comps:
            delta = eff["kpi_delta_t"].get(c["id"], [0, 0])
            comp_vals += [eff["addressable_t"].get(c["id"], 0),
                          delta[0], delta[1]]
        w.writerow([
            h["rank"], h["scores"]["priority"], h["title"], h["category_ru"],
            h["hypothesis"],
            *comp_vals,
            h["scores"]["feasibility"], h["scores"]["novelty"],
            h["scores"]["risk"], h["status"],
            "; ".join(h["streams"]), "; ".join(h["sources"]),
        ])
    return buf.getvalue()


def to_markdown(result: dict) -> str:
    fmt = lambda v: f"{v:,.0f}".replace(",", " ")
    s = result["summary"]
    comps = _components(result)
    ids = [c["id"] for c in comps]
    losses = s.get("losses_t") or {i: s.get(f"losses_{i}_t", 0) for i in ids}
    recov = s.get("recoverable_t") or {i: s.get(f"recoverable_{i}_t", 0) for i in ids}
    rec_pct = s.get("recoverable_pct") or {i: s.get(f"recoverable_{i}_pct", 0) for i in ids}
    lines = [
        f"# Гипотезы снижения потерь металлов — {result['plant']}",
        "",
        f"*Движок генерации: {result['engine']}*",
        "",
    ]
    proj = result.get("project")
    if proj:
        lines += [f"**Задача:** {proj.get('target_kpi') or '—'}", ""]
        if proj.get("constraints"):
            lines += ["**Ограничения:** " + "; ".join(proj["constraints"]), ""]
    lines += [
        "## Сводка потерь",
        "",
        "| Показатель | " + " | ".join(c["label"] for c in comps) + " |",
        "|---|" + "---|" * len(comps),
        "| Потери с хвостами, т | " + " | ".join(fmt(losses.get(i, 0)) for i in ids) + " |",
        "| Извлекаемый металл, т | " + " | ".join(fmt(recov.get(i, 0)) for i in ids) + " |",
        "| Доля извлекаемого, % | " + " | ".join(str(rec_pct.get(i, 0)) for i in ids) + " |",
        "",
        "## Ключевые находки диагностики",
        "",
    ]
    for f_ in result["findings"][:8]:
        prefix = "(справочно) " if f_.get("informational") else ""
        lines.append(f"- {prefix}**{f_['title']}** — {fmt(f_['tons'])} т "
                     f"({f_['share_of_losses_pct']}% потерь потока). {f_['detail']}")
    lines += ["", "## Гипотезы (по приоритету)", ""]

    for h in result["hypotheses"]:
        eff = h["expected_effect"]
        sc = h["scores"]
        lines += [
            f"### {h['rank']}. {h['title']}",
            "",
            f"**Приоритет: {sc['priority']}** · категория: {h['category_ru']} · "
            f"эффект: {fmt(sc['impact_t'])} т адресуемо · реализуемость {sc['feasibility']}/5 · "
            f"новизна {sc['novelty']}/5 · риск {sc['risk']}/5 · статус: {h['status']}",
            "",
            f"**Гипотеза.** {h['hypothesis']}",
            "",
            f"**Механизм.** {h['mechanism']}",
            "",
            "**Обоснование и источники:**",
        ]
        for e in h["evidence"]:
            lines.append(f"- {e['fact']}  \n  *— {e['source']}*")
        delta_line = ", ".join(
            f"{c['label']} −{eff['kpi_delta_t'].get(c['id'], [0, 0])[0]:.0f}…"
            f"{eff['kpi_delta_t'].get(c['id'], [0, 0])[1]:.0f} т" for c in comps)
        lines += [
            "",
            f"**Ожидаемый KPI.** {eff['kpi']}: {delta_line}. {eff['assumption']}",
            "",
            "**Риски:** " + "; ".join(h["risks"]),
            "",
            "**Дорожная карта проверки:**",
        ]
        lines += [f"{i}. {step}" for i, step in enumerate(h["roadmap"], 1)]
        lines.append("")

    return "\n".join(lines)


def to_docx(result: dict) -> bytes:
    """DOCX-отчёт (требование ТЗ «Генерация отчётов PDF/DOCX»)."""
    import docx

    doc = docx.Document()
    doc.add_heading(f"Гипотезы снижения потерь металлов — {result['plant']}", 0)
    doc.add_paragraph(f"Движок генерации: {result['engine']}")
    proj = result.get("project")
    if proj:
        doc.add_paragraph(f"Задача: {proj.get('target_kpi') or '—'}")
        if proj.get("constraints"):
            doc.add_paragraph("Ограничения: " + "; ".join(proj["constraints"]))

    s = result["summary"]
    comps = _components(result)
    ids = [c["id"] for c in comps]
    losses = s.get("losses_t") or {i: s.get(f"losses_{i}_t", 0) for i in ids}
    recov = s.get("recoverable_t") or {i: s.get(f"recoverable_{i}_t", 0) for i in ids}
    rec_pct = s.get("recoverable_pct") or {i: s.get(f"recoverable_{i}_pct", 0) for i in ids}
    doc.add_heading("Сводка потерь", level=1)
    table = doc.add_table(rows=4, cols=1 + len(comps))
    table.style = "Table Grid"
    rows = [
        ("Показатель", *(c["label"] for c in comps)),
        ("Потери с хвостами, т", *(f"{losses.get(i, 0):,.0f}" for i in ids)),
        ("Извлекаемый металл, т", *(f"{recov.get(i, 0):,.0f}" for i in ids)),
        ("Доля извлекаемого, %", *(str(rec_pct.get(i, 0)) for i in ids)),
    ]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val.replace(",", " ")

    doc.add_heading("Ключевые находки диагностики", level=1)
    for f_ in result["findings"][:8]:
        prefix = "(справочно) " if f_.get("informational") else ""
        doc.add_paragraph(
            f"{prefix}{f_['title']} — {f_['tons']:,.0f} т "
            f"({f_['share_of_losses_pct']}% потерь потока). {f_['detail']}"
            .replace(",", " "), style="List Bullet")

    doc.add_heading("Гипотезы (по приоритету)", level=1)
    for h in result["hypotheses"]:
        sc, eff = h["scores"], h["expected_effect"]
        doc.add_heading(f"{h['rank']}. {h['title']}", level=2)
        doc.add_paragraph(
            f"Приоритет {sc['priority']} · {h['category_ru']} · "
            f"эффект {sc['impact_t']:,.0f} т · реализуемость {sc['feasibility']}/5 · "
            f"новизна {sc['novelty']}/5 · риск {sc['risk']}/5 · статус {h['status']}"
            .replace(",", " "))
        doc.add_paragraph(f"Гипотеза. {h['hypothesis']}")
        doc.add_paragraph(f"Механизм. {h['mechanism']}")
        doc.add_paragraph("Обоснование и источники:")
        for e in h["evidence"]:
            doc.add_paragraph(f"{e['fact']} — {e['source']}", style="List Bullet")
        delta_line = ", ".join(
            f"{c['label']} −{eff['kpi_delta_t'].get(c['id'], [0, 0])[0]:.0f}…"
            f"{eff['kpi_delta_t'].get(c['id'], [0, 0])[1]:.0f} т" for c in comps)
        doc.add_paragraph(f"Ожидаемый KPI: {delta_line}. {eff['assumption']}")
        doc.add_paragraph("Риски: " + "; ".join(h["risks"]))
        doc.add_paragraph("Дорожная карта проверки:")
        for i, step in enumerate(h["roadmap"], 1):
            doc.add_paragraph(f"{i}. {step}", style="List Number")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
