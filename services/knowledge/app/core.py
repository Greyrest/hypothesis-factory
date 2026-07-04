"""База знаний для генерации гипотез.

Три источника:
1. Справка института «Как читать отчёт» (docx) -> чанки для retrieval с цитированием.
2. Эталонные гипотезы экспертов из 4 примеров (docx) -> каталог проверенных практик,
   обогащённый курируемыми метаданными (категория, оборудование, оценки).
3. Доменные правила физики обогащения (курируемые, с привязкой к литературе).

Категории гипотез:
GRIND    — измельчение (футеровка, шары, водный режим)
CLASSIFY — классификация (гидроциклоны, классификаторы, грохочение)
REGRIND  — доизмельчение целевого класса в отдельном цикле, магнитная сепарация
FLOT     — флотация (фронт, время, плотность пульпы, контактные чаны)
REAGENT  — реагентный режим
CRUSH    — дробление / питание мельниц
TAILS    — переработка хвостовых потоков (классификация хвостов, возврат)
AUTO     — автоматизация и контроль
"""
from __future__ import annotations

import json
import re
from pathlib import Path

import docx

# --- курируемые метаданные каталога: ключ-regex -> атрибуты -----------------
# feasibility/novelty: 1..5 (больше = лучше/новее); risk: 1..5 (больше = рискованнее)
CATALOG_META = [
    (r"магнитная сепарация.*доизмельчением", dict(
        categories=["REGRIND"], signals=["coarse_locked", "pyrrhotite"],
        equipment="магнитный сепаратор + мельница доизмельчения",
        feasibility=2, risk=3, novelty=3, capex="высокий")),
    (r"футеровки", dict(
        categories=["GRIND"], signals=["coarse_locked"],
        equipment="шаровые мельницы", feasibility=3, risk=2, novelty=2, capex="средний")),
    (r"песковых насадок на гидроциклонах", dict(
        categories=["CLASSIFY"], signals=["coarse_locked", "coarse_share"],
        equipment="гидроциклоны", feasibility=4, risk=2, novelty=2, capex="низкий")),
    (r"замена классификаторов на гидроциклоны", dict(
        categories=["CLASSIFY"], signals=["coarse_locked", "coarse_share"],
        equipment="классификаторы -> гидроциклоны", feasibility=2, risk=3, novelty=2,
        capex="высокий")),
    (r"грохота тонкого грохочения", dict(
        categories=["CLASSIFY", "REGRIND"], signals=["coarse_locked"],
        equipment="грохоты тонкого грохочения", feasibility=2, risk=3, novelty=3,
        capex="высокий")),
    (r"контроль гранулометрии руды", dict(
        categories=["CRUSH", "AUTO"], signals=["coarse_locked", "coarse_share"],
        equipment="конусные дробилки, гранулометры", feasibility=4, risk=1, novelty=2,
        capex="низкий")),
    (r"зазора щели конусных дробилок", dict(
        categories=["CRUSH", "AUTO"], signals=["coarse_locked", "coarse_share"],
        equipment="конусные дробилки", feasibility=3, risk=2, novelty=2, capex="средний")),
    (r"подачи воды в шаровые мельницы", dict(
        categories=["GRIND", "AUTO"], signals=["coarse_locked"],
        equipment="шаровые мельницы, АСУ ТП", feasibility=4, risk=1, novelty=2,
        capex="низкий")),
    (r"насадок на классификаторах", dict(
        categories=["CLASSIFY"], signals=["coarse_locked", "coarse_share"],
        equipment="классификаторы", feasibility=4, risk=2, novelty=2, capex="низкий")),
    (r"фронта флотации|контрольной операции", dict(
        categories=["FLOT"], signals=["mid_liberated", "fine_liberated"],
        equipment="флотомашины", feasibility=3, risk=2, novelty=2, capex="средний")),
    (r"плотности пульпы", dict(
        categories=["FLOT"], signals=["mid_liberated", "fine_liberated"],
        equipment="узел кондиционирования", feasibility=5, risk=2, novelty=2,
        capex="низкий")),
    (r"контактных чанов", dict(
        categories=["FLOT", "REAGENT"], signals=["mid_liberated", "fine_liberated"],
        equipment="контактные чаны", feasibility=3, risk=2, novelty=3, capex="средний")),
    (r"дробилки рудной гали", dict(
        categories=["CRUSH", "GRIND"], signals=["coarse_locked"],
        equipment="дробилка гали", feasibility=2, risk=3, novelty=3, capex="высокий")),
    (r"finfix", dict(
        categories=["REAGENT"], signals=["fine_liberated"],
        equipment="реагентное хозяйство", feasibility=4, risk=3, novelty=4,
        capex="низкий")),
    (r"мелящих шаров диаметром|шаров диаметром 120", dict(
        categories=["GRIND"], signals=["coarse_locked"],
        equipment="шаровые мельницы", feasibility=4, risk=2, novelty=2, capex="низкий")),
    (r"шары 5-го класса|повышенной износостойкости", dict(
        categories=["GRIND"], signals=["coarse_locked"],
        equipment="шаровые мельницы", feasibility=4, risk=2, novelty=3, capex="низкий")),
    (r"скорости вращения классификаторов", dict(
        categories=["CLASSIFY", "AUTO"], signals=["coarse_locked", "coarse_share"],
        equipment="спиральные классификаторы", feasibility=5, risk=1, novelty=2,
        capex="низкий")),
    (r"дополнительная классификация целевого класса", dict(
        categories=["CLASSIFY", "REGRIND"], signals=["coarse_locked"],
        equipment="гидроциклоны / грохоты", feasibility=3, risk=2, novelty=3,
        capex="средний")),
    (r"классификаторов на более производительные", dict(
        categories=["CLASSIFY"], signals=["coarse_share", "coarse_locked"],
        equipment="классификаторы", feasibility=2, risk=3, novelty=2, capex="высокий")),
    (r"эффективности гидроциклонов", dict(
        categories=["CLASSIFY"], signals=["coarse_share", "coarse_locked"],
        equipment="гидроциклоны", feasibility=4, risk=1, novelty=2, capex="низкий")),
    (r"классификация хвостов и возврат", dict(
        categories=["TAILS"], signals=["coarse_locked", "mid_liberated"],
        equipment="узел классификации хвостов", feasibility=2, risk=3, novelty=4,
        capex="высокий")),
    (r"контрольная классификация", dict(
        categories=["CLASSIFY"], signals=["coarse_share", "coarse_locked"],
        equipment="классификаторы", feasibility=3, risk=2, novelty=3, capex="средний")),
]

DEFAULT_META = dict(categories=["CLASSIFY"], signals=["coarse_locked"],
                    equipment="", feasibility=3, risk=2, novelty=2, capex="средний")

# --- доменные правила (сигнал -> физика -> категории решений) ---------------
DOMAIN_RULES = [
    dict(id="R1", signal="coarse_locked",
         title="Закрытый металл в крупных классах — недоизмельчение",
         text="Закрытый Pnt/Cp — металл в сростках с породой. Его преобладание в "
              "крупных классах (+45 мкм и крупнее) означает, что руда не домолота: "
              "минерал не раскрыт и не может быть сфлотирован. Решения — повышение "
              "тонины помола (измельчение) и исключение проскока крупных частиц "
              "(классификация), либо доизмельчение выделенного класса в отдельном "
              "цикле.",
         categories=["GRIND", "CLASSIFY", "REGRIND", "CRUSH"],
         source="Справка «Как читать отчёт»; Абрамов А.А. «Флотационные методы обогащения» (раскрытие минералов)"),
    dict(id="R2", signal="fine_liberated",
         title="Раскрытый металл в классе -10 мкм — шламовые потери флотации",
         text="Раскрытый (свободный) минерал в тоне -10 мкм флотация теряет из-за "
              "малой массы частиц: низкая вероятность закрепления на пузырьке, "
              "вынос со сливом. Переизмельчение усугубляет проблему. Решения — "
              "реагентный режим (флокулянты/собиратели для шламов), время "
              "флотации, плотность пульпы; а также не переизмельчать (контроль "
              "циркулирующей нагрузки).",
         categories=["REAGENT", "FLOT", "GRIND"],
         source="Абрамов А.А. «Флотационные методы обогащения», гл. о флотации шламов"),
    dict(id="R3", signal="mid_liberated",
         title="Раскрытый металл в средних классах — недостаток флотационной ёмкости",
         text="Свободный минерал флотационной крупности (-71..-10 мкм) в хвостах — "
              "признак нехватки времени флотации, фронта машин или неоптимальной "
              "плотности/аэрации. Такой металл извлекается настройкой режима без "
              "капитальных затрат.",
         categories=["FLOT", "REAGENT"],
         source="Справка «Как читать отчёт»; практика НОФ (Пример 3)"),
    dict(id="R4", signal="mid_locked",
         title="Сростки в средних классах — доизмельчение промпродукта",
         text="Закрытый металл в классах -71+20 мкм говорит о неполном раскрытии "
              "при штатной тонине помола. Целесообразно доизмельчение "
              "промпродукта/песков контрольной классификации.",
         categories=["REGRIND", "GRIND"],
         source="Технология обогащения полезных ископаемых (geokniga), раздел «Измельчение»"),
    dict(id="R5", signal="coarse_share",
         title="Высокая доля крупных классов в хвостах — проскок классификации",
         text="Если заметная доля массы хвостов крупнее 45–71 мкм, классификация "
              "пропускает крупные частицы в флотацию и далее в отвал. Проверить "
              "насадки/зазоры, циркулирующую нагрузку, эффективность "
              "гидроциклонов/классификаторов.",
         categories=["CLASSIFY", "CRUSH", "AUTO"],
         source="Справка «Как читать отчёт»; регламенты фабрик"),
    dict(id="R6", signal="pyrrhotite",
         title="Металл, связанный с пирротином",
         text="Примесь в пирротине текущей технологией не извлекается (изоморфная "
              "примесь). Однако для пирротиновых хвостов извлекаемая часть "
              "(раскрытые/закрытые сульфиды) может выделяться магнитной "
              "сепарацией или доизвлекаться в отдельном цикле.",
         categories=["REGRIND", "TAILS"],
         source="Справка «Как читать отчёт»; практика КГМК (Пример 1)"),
    dict(id="R7", signal="tails_recycle",
         title="Возврат песковой части хвостов",
         text="Если извлекаемый металл сконцентрирован в узком классе хвостов, "
              "экономически оправдана классификация хвостов с возвратом песковой "
              "части в голову процесса.",
         categories=["TAILS"],
         source="Практика ТОФ (Пример 4)"),
]


def _doc_meta(path: Path) -> dict:
    """Метаданные документа-источника (ТЗ: источники, даты, авторы)."""
    try:
        props = docx.Document(str(path)).core_properties
        return {
            "author": props.author or None,
            "date": str(props.created.date()) if props.created else None,
            "file": path.name,
        }
    except Exception:
        return {"author": None, "date": None, "file": path.name}


def _read_docx_paragraphs(path: Path) -> list[str]:
    d = docx.Document(str(path))
    out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            seen = set()
            for c in row.cells:
                t = c.text.strip()
                if t and t not in seen:
                    out.append(t)
                    seen.add(t)
    return out


def _catalog_meta(title: str) -> dict:
    low = title.lower()
    for pattern, meta in CATALOG_META:
        if re.search(pattern, low):
            return dict(meta)
    return dict(DEFAULT_META)


def build_kb(data_dir: str | Path) -> dict:
    data_dir = Path(data_dir)
    kb = {"chunks": [], "catalog": [], "rules": DOMAIN_RULES}

    # 1. справка
    guide = next(data_dir.glob("Как читать отчет*.docx"), None)
    if guide:
        gmeta = _doc_meta(guide)
        for i, p in enumerate(_read_docx_paragraphs(guide)):
            kb["chunks"].append({
                "id": f"guide-{i}", "kind": "guide",
                "source": "Справка института «Как читать отчёт по хвостам»",
                "meta": gmeta,
                "text": p})

    # 2. каталог экспертных гипотез из примеров
    seen_titles = set()
    for hyp_file in sorted(data_dir.glob("Пример */Гипотезы*.docx")):
        plant = hyp_file.parent.name  # 'Пример 1'
        label = hyp_file.stem.replace("Гипотезы", "").strip()
        doc_meta = _doc_meta(hyp_file)
        for line in _read_docx_paragraphs(hyp_file):
            m = re.match(r"^\d+\.\s*(.+)$", line)
            if not m:
                continue
            title = m.group(1).strip().rstrip(".")
            key = re.sub(r"\W+", "", title.lower())[:60]
            meta = _catalog_meta(title)
            entry = {
                "id": f"cat-{len(kb['catalog'])+1:02d}",
                "title": title,
                "source": f"Мозговой штурм экспертов, {plant} ({label})",
                "meta": doc_meta,
                "plants": [label],
                **meta,
            }
            if key in seen_titles:
                # дубль между фабриками -> расширяем список фабрик
                for e in kb["catalog"]:
                    if re.sub(r"\W+", "", e["title"].lower())[:60] == key:
                        e["plants"].append(label)
                        e["source"] += f"; {plant} ({label})"
                continue
            seen_titles.add(key)
            kb["catalog"].append(entry)

    # 3. правила как чанки (для retrieval)
    for r in DOMAIN_RULES:
        kb["chunks"].append({
            "id": r["id"], "kind": "rule", "source": r["source"],
            "text": f"{r['title']}. {r['text']}"})

    # 4. литература (метазаписи — на что ссылаться в дорожных картах)
    extras = data_dir / "Дополнительные материалы"
    if extras.exists():
        for pdf in sorted(extras.glob("*.pdf")):
            kb["chunks"].append({
                "id": f"lit-{pdf.stem[:30]}", "kind": "literature",
                "source": pdf.name,
                "text": f"Учебник/монография: {pdf.stem}"})

    return kb


def retrieve(kb: dict, query_terms: list[str], kinds=("guide", "rule"), top_k=6) -> list[dict]:
    """Простой лексический retrieval: скоринг по пересечению термов."""
    q = {t.lower() for t in query_terms}
    scored = []
    for ch in kb["chunks"]:
        if ch["kind"] not in kinds:
            continue
        words = set(re.findall(r"[а-яёa-z0-9+-]+", ch["text"].lower()))
        score = len(q & words)
        # сигнальные id правил дают бонус
        if ch["kind"] == "rule" and any(t in ch["text"].lower() for t in q):
            score += 2
        if score:
            scored.append((score, ch))
    scored.sort(key=lambda x: -x[0])
    return [ch for _, ch in scored[:top_k]]


if __name__ == "__main__":
    import sys
    kb = build_kb(sys.argv[1] if len(sys.argv) > 1 else
                  "Задача 1. Фабрика гипотез/Задача 1")
    print(json.dumps({"chunks": len(kb["chunks"]),
                      "catalog": len(kb["catalog"]),
                      "rules": len(kb["rules"])}, ensure_ascii=False))
    for e in kb["catalog"]:
        print(f"  {e['id']} [{','.join(e['categories'])}] {e['title'][:70]} <- {', '.join(e['plants'])}")
